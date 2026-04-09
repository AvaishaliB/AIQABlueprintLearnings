"""
Sanity Check Report Generator
Creates a Word document with embedded screenshot from SauceDemo login test
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from datetime import datetime

def create_sanity_report(screenshot_path, output_filename="Sanity_Result.docx"):
    """
    Creates a Word document with the sanity check screenshot
    
    Args:
        screenshot_path (str): Path to the screenshot file
        output_filename (str): Name of the output docx file
    """
    
    # Check if screenshot exists
    if not os.path.exists(screenshot_path):
        print(f"Error: Screenshot file not found at {screenshot_path}")
        print("Please save your screenshot first and provide the correct path.")
        return False
    
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading("Sanity Check Report", 0)
        title.alignment = 1  # Center alignment
        
        # Add test details
        doc.add_heading("Test Details", level=1)
        details = doc.add_paragraph()
        details.add_run("Application: ").bold = True
        details.add_run("SauceDemo (https://www.saucedemo.com/)\n")
        
        details.add_run("Test Type: ").bold = True
        details.add_run("Login Sanity Check\n")
        
        details.add_run("Username: ").bold = True
        details.add_run("standard_user\n")
        
        details.add_run("Timestamp: ").bold = True
        details.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        
        # Add test result heading
        doc.add_heading("Test Result", level=1)
        doc.add_paragraph("Login successful. Screenshot showing post-login dashboard:")
        
        # Add screenshot
        doc.add_heading("Screenshot", level=2)
        doc.add_picture(screenshot_path, width=Inches(6))
        
        # Add footer
        doc.add_paragraph("\n")
        footer = doc.add_paragraph("Report generated automatically by Sanity Check Script")
        footer.runs[0].font.size = Pt(10)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Save document
        doc.save(output_filename)
        print(f"✅ Success! Report saved as: {output_filename}")
        print(f"📁 Location: {os.path.abspath(output_filename)}")
        return True
        
    except Exception as e:
        print(f"❌ Error creating document: {str(e)}")
        return False

if __name__ == "__main__":
    print("=== Sanity Check Report Generator ===\n")
    
    # Get screenshot path from user
    screenshot_path = input("Enter the path to your screenshot file\n(e.g., C:\\Users\\YourName\\Pictures\\screenshot.png): ").strip()
    
    if not screenshot_path:
        print("No path provided. Using default: screenshot.png")
        screenshot_path = "screenshot.png"
    
    # Create the report
    success = create_sanity_report(screenshot_path)
    
    if success:
        print("\n✨ Ready to use! Open Sanity_Result.docx to view your report.")
    else:
        print("\n⚠️  Please ensure your screenshot file exists and try again.")
